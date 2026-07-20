"""PDF, Word and Excel scientific reports built from persisted artifacts."""

from __future__ import annotations

import json
from pathlib import Path

import pandas as pd


REPORT_LABELS = {
    "model": "Modelo", "model_a": "Modelo A", "model_b": "Modelo B", "fold": "Pliegue", "accuracy": "Exactitud",
    "precision": "Precisión", "recall": "Sensibilidad", "specificity": "Especificidad", "f1": "F1",
    "roc_auc": "ROC-AUC", "pr_auc": "PR-AUC", "log_loss": "Pérdida logarítmica", "mcc": "Coeficiente de Matthews", "brier_score": "Puntuación de Brier",
    "threshold": "Umbral", "parameters": "Parámetros", "artifact_bytes": "Tamaño del artefacto (bytes)", "bootstrap_f1": "IC bootstrap de F1", "bootstrap_roc_auc": "IC bootstrap de ROC-AUC",
    "balanced_accuracy": "Exactitud balanceada", "training_seconds": "Entrenamiento (s)",
    "inference_seconds_per_row": "Inferencia por fila (s)", "inference_p50_seconds": "Inferencia P50 (s)",
    "inference_p95_seconds": "Inferencia P95 (s)", "cv_f1_mean": "F1 medio CV", "cv_f1_std": "Desviación F1 CV",
    "cv_roc_auc_mean": "ROC-AUC medio CV", "cv_recall_mean": "Sensibilidad media CV", "p_value": "Valor p",
    "p_value_holm": "Valor p (Holm)", "mean_difference": "Diferencia media", "sd_difference": "Desviación de la diferencia",
    "degrees_of_freedom": "Grados de libertad", "ci95_low": "IC95% inferior", "ci95_high": "IC95% superior",
    "t_statistic": "Estadístico t", "cohen_d_paired": "d de Cohen pareado", "n": "n", "n_nonzero": "n no nulo",
    "positive_rank_sum": "Suma de rangos positivos", "negative_rank_sum": "Suma de rangos negativos", "statistic": "Estadístico",
    "effect": "Efecto rango-biserial", "kendall_w": "W de Kendall", "significant": "Significativo", "probability": "Probabilidad",
    "temperature": "Temperatura", "humidity": "Humedad", "moi": "MOI", "soil_type": "Tipo de suelo", "crop_stage": "Etapa del cultivo", "crop_id": "Identificador del cultivo", "result": "Resultado",
    "predicted_class": "Clase predicha", "class": "Clase", "recommendation": "Recomendación", "timestamp": "Fecha y hora",
    "input": "Entrada", "model_name": "Nombre del modelo", "model_version": "Versión del modelo", "requested_folds": "Pliegues solicitados",
    "fast_mode": "Modo rápido", "epochs": "Épocas", "feature_columns": "Columnas de entrada", "artifact_files": "Archivos generados",
    "seed": "Semilla", "platform": "Plataforma", "rows": "Registros", "rows_excluded": "Filas excluidas", "columns": "Columnas",
    "duplicates": "Duplicados", "missing_by_column": "Faltantes por columna", "dtypes": "Tipos de datos",
    "class_distribution": "Distribución de clases", "numeric_describe": "Estadísticos descriptivos", "warnings": "Advertencias",
    "normalized_columns": "Columnas normalizadas", "cleaned_dataset": "Conjunto de datos limpio", "value": "Valor", "best_value": "Mejor valor",
    "best_params": "Mejores parámetros", "backend": "Motor", "trials": "Pruebas", "interpretation_es": "Interpretación",
    "null_hypothesis_es": "Hipótesis nula", "methodology_es": "Metodología", "paired_t": "t pareada", "wilcoxon": "Wilcoxon",
    "mcnemar": "McNemar", "n_folds": "Número de pliegues", "available": "Disponible", "reason": "Motivo", "direction": "Dirección",
    "models": "Modelos", "alpha": "Alfa", "friedman": "Friedman", "pairwise": "Comparaciones por pares", "message": "Mensaje",
    "folds": "Pliegues", "metrics": "Métricas", "tuning": "Ajuste de hiperparámetros", "statistical_tests": "Pruebas estadísticas",
    "count": "Cantidad", "unique": "Valores únicos", "top": "Más frecuente", "freq": "Frecuencia", "mean": "Media", "std": "Desviación estándar",
    "min": "Mínimo", "max": "Máximo", "25%": "Percentil 25", "50%": "Mediana", "75%": "Percentil 75",
}


def _spanish_columns(frame: pd.DataFrame) -> pd.DataFrame:
    return frame.rename(columns={column: REPORT_LABELS.get(column, column) for column in frame.columns})


def _spanish_payload(value):
    if isinstance(value, dict):
        return {REPORT_LABELS.get(key, key): _spanish_payload(item) for key, item in value.items()}
    if isinstance(value, list):
        return [_spanish_payload(item) for item in value]
    return value


def _figure_title(path: Path) -> str:
    titles = {
        "confusion_matrix_best": "Matriz de confusión del mejor modelo", "roc_curves_models": "Curvas ROC de los modelos",
        "model_metrics_heatmap": "Mapa de calor de métricas por modelo", "correlation_heatmap": "Mapa de calor de correlaciones",
        "distribution_soil_type": "Distribución por tipo de suelo", "distribution_crop_stage": "Distribución por etapa del cultivo",
        "numeric_boxplots": "Diagramas de caja de variables numéricas", "numeric_histograms": "Histogramas de variables numéricas",
    }
    return titles.get(path.stem, path.stem.replace("_", " ").capitalize())


def _figures(artifacts: Path) -> list[Path]:
    return sorted((artifacts / "figures").glob("*.png"))


def _load_stats(artifacts: Path) -> dict:
    path = artifacts / "metrics" / "statistical_tests.json"
    return json.loads(path.read_text(encoding="utf-8")) if path.exists() else {}


def _interpretation(stats: dict) -> str:
    friedman = stats.get("friedman") or {}
    if friedman.get("available") is False:
        return "Friedman requiere al menos tres folds completos; ejecute el modo científico con 5 folds para obtener el contraste global."
    p_value = friedman.get("p_value")
    if p_value is None:
        return "No fue posible calcular Friedman con los folds disponibles."
    global_text = (
        "Friedman detecta diferencias globales entre los modelos; el resultado no identifica por sí solo qué pares difieren."
        if float(p_value) < 0.05 else
        "Friedman no detecta diferencias globales estadísticamente significativas con alfa=0.05."
    )
    paired = [row for row in stats.get("paired_t", []) if row.get("interpretation_es")]
    wilcoxon = [row for row in stats.get("wilcoxon", []) if row.get("interpretation_es")]
    return f"{global_text} t pareada: {sum(bool(row.get('significant')) for row in paired)}/{len(paired)} comparaciones significativas. Wilcoxon: {sum(bool(row.get('significant')) for row in wilcoxon)}/{len(wilcoxon)} significativas tras Holm."


def generate_reports(artifacts: Path | str, frame: pd.DataFrame, comparison: pd.DataFrame, cv: pd.DataFrame) -> None:
    artifacts = Path(artifacts)
    reports = artifacts / "reports"
    reports.mkdir(parents=True, exist_ok=True)
    metadata_path = artifacts / "models" / "model_metadata.json"
    metadata = json.loads(metadata_path.read_text(encoding="utf-8")) if metadata_path.exists() else {}
    stats = _load_stats(artifacts)
    tuning_path = artifacts / "metrics" / "tuning.json"
    tuning = json.loads(tuning_path.read_text(encoding="utf-8")) if tuning_path.exists() else {}
    quality_path = artifacts / "metrics" / "data_quality.json"
    quality = json.loads(quality_path.read_text(encoding="utf-8")) if quality_path.exists() else {}

    # Excel: one sheet per evidence family so the user can filter and audit it.
    with pd.ExcelWriter(reports / "smart_irrigation_report.xlsx", engine="openpyxl") as writer:
        _spanish_columns(frame.head(1000)).to_excel(writer, sheet_name="Datos_muestra", index=False)
        _spanish_columns(frame.describe(include="all").transpose()).to_excel(writer, sheet_name="EDA_descriptivo")
        _spanish_columns(pd.DataFrame([quality])).to_excel(writer, sheet_name="Calidad", index=False)
        _spanish_columns(comparison).to_excel(writer, sheet_name="Metricas_modelos", index=False)
        _spanish_columns(cv).to_excel(writer, sheet_name="Validacion_cruzada", index=False)
        _spanish_columns(pd.DataFrame(tuning).T).to_excel(writer, sheet_name="Hiperparametros")
        _spanish_columns(pd.DataFrame(stats.get("paired_t", []))).to_excel(writer, sheet_name="t_pareada", index=False)
        _spanish_columns(pd.DataFrame(stats.get("wilcoxon", stats.get("pairwise", [])))).to_excel(writer, sheet_name="Wilcoxon_Holm", index=False)
        _spanish_columns(pd.DataFrame([stats.get("friedman", {})])).to_excel(writer, sheet_name="Friedman", index=False)
        _spanish_columns(pd.DataFrame(stats.get("mcnemar", []))).to_excel(writer, sheet_name="McNemar", index=False)
        _spanish_columns(pd.DataFrame(columns=["timestamp", "input", "probability", "class", "recommendation"])).to_excel(writer, sheet_name="Predicciones", index=False)
        _spanish_columns(pd.DataFrame([metadata])).to_excel(writer, sheet_name="Modelo_guardado", index=False)

    # Word: readable narrative plus tables and all generated figures.
    from docx import Document
    from docx.shared import Inches

    doc = Document()
    doc.add_heading("Sistema de riego inteligente", 0)
    doc.add_paragraph("Informe científico generado automáticamente a partir del conjunto de datos cargado manualmente y de la ejecución real del flujo.")
    doc.add_heading("Resumen ejecutivo", level=1)
    doc.add_paragraph(f"Se analizaron {len(frame):,} registros útiles. El mejor modelo persistido es {metadata.get('model_name', 'pendiente')}. Se excluyeron {quality.get('rows_excluded', 0):,} filas con result=2 por corresponder a exceso de agua en el estudio binario.")
    doc.add_paragraph(_interpretation(stats))
    doc.add_heading("Modelo y configuración", level=1)
    for key in ["model_name", "model_version", "folds", "requested_folds", "fast_mode", "epochs", "feature_columns"]:
        if key in metadata:
            doc.add_paragraph(f"{REPORT_LABELS.get(key, key)}: {metadata[key]}")
    doc.add_heading("Comparación de modelos", level=1)
    columns = [c for c in ["model", "accuracy", "precision", "recall", "f1", "roc_auc", "training_seconds", "inference_p95_seconds"] if c in comparison.columns]
    table = doc.add_table(rows=1, cols=len(columns))
    for cell, name in zip(table.rows[0].cells, columns):
        cell.text = REPORT_LABELS.get(name, name)
    for _, row in comparison[columns].iterrows():
        cells = table.add_row().cells
        for cell, name in zip(cells, columns):
            cell.text = str(row[name])
    doc.add_heading("Pruebas estadísticas", level=1)
    doc.add_paragraph("t de Student pareada: compara la diferencia media de F1 entre modelos en los mismos pliegues. Wilcoxon: alternativa no paramétrica pareada. Friedman: contraste global para tres o más modelos. Holm controla el error por comparaciones múltiples.")
    doc.add_paragraph(_interpretation(stats))
    for title, key in [("Friedman", "friedman"), ("t de Student pareada", "paired_t"), ("Wilcoxon con Holm", "wilcoxon")]:
        doc.add_heading(title, level=2)
        value = stats.get(key, {})
        doc.add_paragraph(json.dumps(_spanish_payload(value), ensure_ascii=False, indent=2, default=str)[:5000])
        if isinstance(value, list):
            for row in value:
                if row.get("interpretation_es"):
                    doc.add_paragraph(row["interpretation_es"], style="List Bullet")
    doc.add_heading("Figuras", level=1)
    for figure in _figures(artifacts):
        doc.add_paragraph(_figure_title(figure))
        doc.add_picture(str(figure), width=Inches(5.8))
    doc.add_heading("Limitaciones y uso responsable", level=1)
    doc.add_paragraph("Las arquitecturas CNN y LSTM reciben el vector tabular como una secuencia compacta; esto no equivale a disponer de una serie temporal real. La predicción es una recomendación de apoyo y debe contrastarse con las condiciones agronómicas del cultivo. No se declara superioridad si el p-valor ajustado es mayor o igual que 0.05.")
    doc.save(reports / "smart_irrigation_report.docx")

    # PDF: compact executive version with the same evidence and interpretation.
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import Image, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

    styles = getSampleStyleSheet()
    story = [
        Paragraph("Sistema de riego inteligente", styles["Title"]),
        Paragraph("Informe científico generado con resultados reales.", styles["Normal"]),
        Spacer(1, 0.18 * inch),
        Paragraph(f"Registros útiles: {len(frame):,}. Filas excluidas: {quality.get('rows_excluded', 0):,}. Modelo seleccionado: {metadata.get('model_name', 'pendiente')}.", styles["Normal"]),
        Paragraph(_interpretation(stats), styles["Normal"]),
        Paragraph("Método: t pareada sobre F1 por pliegue; Friedman por rangos de los modelos; Wilcoxon de rangos con signo sobre diferencias pareadas; Holm corrige las comparaciones múltiples.", styles["Normal"]),
        Spacer(1, 0.14 * inch),
    ]
    pdf_columns = columns[:7]
    table_data = [[REPORT_LABELS.get(column, column) for column in pdf_columns]] + [[str(row[column]) for column in pdf_columns] for _, row in comparison.iterrows()]
    table = Table(table_data, repeatRows=1)
    table.setStyle(TableStyle([("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#075543")), ("TEXTCOLOR", (0, 0), (-1, 0), colors.white), ("GRID", (0, 0), (-1, -1), 0.25, colors.grey), ("FONTSIZE", (0, 0), (-1, -1), 6)]))
    story.extend([table, Spacer(1, 0.18 * inch), Paragraph("Pruebas incluidas: t de Student pareada, Friedman, Wilcoxon con Holm y McNemar cuando los datos lo permiten.", styles["Normal"])])
    for row in stats.get("paired_t", []) + stats.get("wilcoxon", []):
        if row.get("interpretation_es"):
            story.append(Paragraph(str(row["interpretation_es"]), styles["Normal"]))
    for figure in _figures(artifacts):
        story.extend([Spacer(1, 0.1 * inch), Image(str(figure), width=6.3 * inch, height=3.5 * inch)])
    SimpleDocTemplate(str(reports / "smart_irrigation_report.pdf"), pagesize=letter).build(story)
